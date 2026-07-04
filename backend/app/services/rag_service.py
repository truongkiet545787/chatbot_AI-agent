import os
from typing import Dict, Optional, List
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument

def read_pdf(file_path: str) -> str:
    """Đọc và trích xuất toàn bộ văn bản từ file PDF."""
    reader = PdfReader(file_path)
    text_content = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_content.append(text)
    return "\n\n".join(text_content)

def read_docx(file_path: str) -> str:
    """Đọc và trích xuất văn bản từ file Word (.docx), giữ cấu trúc bảng ở dạng Markdown Table."""
    doc = DocxDocument(file_path)
    full_text = []
    
    # Duyệt qua các phần tử trong body theo thứ tự xuất hiện
    for child in doc.element.body:
        if child.tag.endswith('p'):
            # Xử lý đoạn văn
            for p in doc.paragraphs:
                if p._element is child:
                    if p.text.strip():
                        full_text.append(p.text)
                    break
        elif child.tag.endswith('tbl'):
            # Xử lý bảng biểu
            for t in doc.tables:
                if t._element is child:
                    table_lines = []
                    for r_idx, row in enumerate(t.rows):
                        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
                        table_lines.append("| " + " | ".join(cells) + " |")
                        if r_idx == 0:
                            table_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                    if table_lines:
                        full_text.append("\n".join(table_lines))
                    break
                    
    # Dự phòng: Nếu cách duyệt theo thứ tự bị lỗi hoặc trống, quét toàn bộ
    if not full_text:
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for t in doc.tables:
            table_lines = []
            for r_idx, row in enumerate(t.rows):
                cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
                table_lines.append("| " + " | ".join(cells) + " |")
                if r_idx == 0:
                    table_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            if table_lines:
                full_text.append("\n".join(table_lines))
                
    return "\n\n".join(full_text)

class RAGService:
    def __init__(self):
        # Tên mô hình BERT tiếng Việt từ Hugging Face
        model_name = os.getenv("RAG_EMBEDDING_MODEL", "keepitreal/vietnamese-sbert")
        print(f"Initializing RAG Embedding Model: {model_name}...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': 'cpu'}
        )
        # Bộ lưu trữ Chroma in-memory cho từng session_id
        self.stores: Dict[str, Chroma] = {}
        # Danh sách tên các file đã tải lên theo từng session_id
        self.session_files: Dict[str, List[str]] = {}

    def is_session_active(self, session_id: str) -> bool:
        return session_id in self.stores

    def get_session_files(self, session_id: str) -> List[str]:
        return self.session_files.get(session_id, [])

    def add_document(self, session_id: str, file_path: str, filename: str) -> int:
        """
        Đọc tài liệu, chia nhỏ thành chunks, sinh embeddings và lưu vào Chroma in-memory.
        Trả về số lượng chunks được thêm vào.
        """
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            text = read_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            text = read_docx(file_path)
        else:
            raise ValueError(f"Định dạng {ext} không hỗ trợ. Chỉ nhận file PDF hoặc Word (.docx).")

        if not text.strip():
            raise ValueError("Không thể trích xuất văn bản từ tài liệu này.")

        # Phân mảnh tài liệu theo cơ chế Phân cấp Cha-Con (Parent-Child)
        # Đoạn cha (Parent): Chứa trọn vẹn ngữ cảnh của từng phần lớn
        parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=250
        )
        parent_chunks = parent_splitter.split_text(text)
        
        child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,
            chunk_overlap=50
        )
        
        docs = []
        for p_chunk in parent_chunks:
            c_chunks = child_splitter.split_text(p_chunk)
            for c_chunk in c_chunks:
                docs.append(
                    LCDocument(
                        page_content=c_chunk,
                        metadata={
                            "source": filename,
                            "session_id": session_id,
                            "parent_content": p_chunk
                        }
                    )
                )

        # Khởi tạo hoặc cập nhật Chroma vectorstore lưu tạm thời trên RAM
        if session_id in self.stores:
            self.stores[session_id].add_documents(docs)
            if filename not in self.session_files[session_id]:
                self.session_files[session_id].append(filename)
        else:
            vectorstore = Chroma.from_documents(
                documents=docs,
                embedding=self.embeddings,
                collection_name=session_id
            )
            self.stores[session_id] = vectorstore
            self.session_files[session_id] = [filename]

        return len(docs)

    def query_context(self, session_id: str, query: str, k: int = 4) -> Optional[str]:
        """Truy xuất các đoạn văn bản tương đồng ngữ nghĩa làm Context."""
        if session_id not in self.stores:
            return None
            
        vectorstore = self.stores[session_id]
        # Tìm các đoạn con có mức độ tương đồng cao nhất
        docs = vectorstore.similarity_search(query, k=k)
        
        # Trích xuất các đoạn cha tương ứng để chuyển cho LLM
        retrieved_parents = []
        seen_parents = set()
        for doc in docs:
            parent = doc.metadata.get("parent_content")
            if parent:
                parent_strip = parent.strip()
                if parent_strip not in seen_parents:
                    seen_parents.add(parent_strip)
                    retrieved_parents.append(parent)
            else:
                if doc.page_content.strip() not in seen_parents:
                    seen_parents.add(doc.page_content.strip())
                    retrieved_parents.append(doc.page_content)
        
        formatted_context = "\n\n".join(
            f"[Nguồn: {docs[0].metadata.get('source', 'Tài liệu')}]\n{parent_text}"
            for parent_text in retrieved_parents
        )
        return formatted_context

    def summarize_document(self, session_id: str, llm_model) -> Optional[str]:
        """Tóm tắt toàn bộ các tài liệu của session bằng phương pháp Map Reduce thủ công."""
        if session_id not in self.stores:
            return None

        # Lấy tất cả văn bản đã lưu trong Chroma
        store_data = self.stores[session_id].get()
        documents = store_data.get('documents', [])
        
        if not documents:
            return "Không tìm thấy nội dung tài liệu để tóm tắt."

        total_text = "\n\n".join(documents)
        
        # Nếu tài liệu ngắn (< 12.000 ký tự ~ 2500 từ), sử dụng Stuff (Tóm tắt trực tiếp)
        if len(total_text) < 12000:
            prompt = (
                "Bạn là trợ lý phân tích tài liệu chuyên nghiệp. Hãy đọc kỹ tài liệu dưới đây và viết một bản tóm tắt chi tiết, "
                "mạch lạc bằng tiếng Việt, bao gồm các chủ đề chính và các ý quan trọng:\n\n"
                f"Tài liệu:\n{total_text}\n\n"
                "Bản tóm tắt hoàn chỉnh:"
            )
            response = llm_model.invoke(prompt)
            return response.content
        else:
            # Map Reduce: Chia nhỏ văn bản thành các block lớn (~8000 ký tự) để tóm tắt song song/tuần tự
            chunk_summaries = []
            current_block = []
            current_len = 0
            for doc in documents:
                current_block.append(doc)
                current_len += len(doc)
                if current_len >= 8000:
                    block_text = "\n\n".join(current_block)
                    prompt = (
                        "Hãy tóm tắt ngắn gọn các ý chính của phần tài liệu dưới đây bằng tiếng Việt:\n\n"
                        f"{block_text}\n\nTóm tắt ngắn gọn:"
                    )
                    res = llm_model.invoke(prompt)
                    chunk_summaries.append(res.content)
                    current_block = []
                    current_len = 0
            
            if current_block:
                block_text = "\n\n".join(current_block)
                prompt = (
                    "Hãy tóm tắt ngắn gọn các ý chính của phần tài liệu dưới đây bằng tiếng Việt:\n\n"
                    f"{block_text}\n\nTóm tắt ngắn gọn:"
                )
                res = llm_model.invoke(prompt)
                chunk_summaries.append(res.content)

            # Giai đoạn Reduce: Gộp các tóm tắt thành bản tổng hợp cuối cùng
            combined_summaries = "\n\n".join(chunk_summaries)
            final_prompt = (
                "Dưới đây là tóm tắt các phần của một tài liệu lớn. Hãy kết hợp chúng lại thành một bản tóm tắt "
                "hệ thống, chi tiết và mạch lạc bằng tiếng Việt:\n\n"
                f"{combined_summaries}\n\n"
                "Bản tóm tắt hoàn chỉnh:"
            )
            final_res = llm_model.invoke(final_prompt)
            return final_res.content

    def clear_session(self, session_id: str):
        """Xóa sạch dữ liệu RAG của session khỏi RAM."""
        if session_id in self.stores:
            del self.stores[session_id]
        if session_id in self.session_files:
            del self.session_files[session_id]
